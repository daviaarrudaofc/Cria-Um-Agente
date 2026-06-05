from __future__ import annotations

import csv
import os
import re
import sqlite3
from pathlib import Path
from typing import Iterable

from dotenv import load_dotenv
from docx import Document as DocxDocument
from langchain.agents import create_agent
from langchain.tools import tool
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parent
DATA_DIR = ROOT / "data"
INBOX_DIR = DATA_DIR / "inbox"
OUTPUT_DIR = DATA_DIR / "output"
VECTOR_DIR = DATA_DIR / "vectorstore"


def resolve_path(value: str, default: Path) -> Path:
    if not value:
        return default
    path = Path(value)
    return path if path.is_absolute() else ROOT / path


load_dotenv(ROOT / ".env")

DATABASE_PATH = resolve_path(os.getenv("DATABASE_PATH", ""), DATA_DIR / "agent.db")
MODEL_PROVIDER = os.getenv("MODEL_PROVIDER", "xai").strip().lower()
XAI_BASE_URL = os.getenv("XAI_BASE_URL", "https://api.x.ai/v1")
XAI_MODEL = os.getenv("XAI_MODEL", "grok-4.3")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5-mini")
OPENAI_EMBEDDING_MODEL = os.getenv("OPENAI_EMBEDDING_MODEL", "text-embedding-3-small")
DOCUMENT_MEMORY_MODE = os.getenv("DOCUMENT_MEMORY_MODE", "keyword").strip().lower()


def ensure_directories() -> None:
    for directory in (DATA_DIR, INBOX_DIR, OUTPUT_DIR, VECTOR_DIR):
        directory.mkdir(parents=True, exist_ok=True)


def init_sample_database() -> None:
    DATABASE_PATH.parent.mkdir(parents=True, exist_ok=True)
    with sqlite3.connect(DATABASE_PATH) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS clientes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                nome TEXT NOT NULL,
                email TEXT NOT NULL,
                status TEXT NOT NULL
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS tarefas (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                titulo TEXT NOT NULL,
                responsavel TEXT NOT NULL,
                prioridade TEXT NOT NULL,
                concluida INTEGER NOT NULL DEFAULT 0
            )
            """
        )
        if conn.execute("SELECT COUNT(*) FROM clientes").fetchone()[0] == 0:
            conn.executemany(
                "INSERT INTO clientes (nome, email, status) VALUES (?, ?, ?)",
                [
                    ("Ana Souza", "ana@example.com", "ativo"),
                    ("Bruno Lima", "bruno@example.com", "prospect"),
                    ("Carla Dias", "carla@example.com", "ativo"),
                ],
            )
        if conn.execute("SELECT COUNT(*) FROM tarefas").fetchone()[0] == 0:
            conn.executemany(
                """
                INSERT INTO tarefas (titulo, responsavel, prioridade, concluida)
                VALUES (?, ?, ?, ?)
                """,
                [
                    ("Enviar proposta comercial", "Ana Souza", "alta", 0),
                    ("Revisar contrato", "Carla Dias", "media", 0),
                    ("Atualizar planilha de leads", "Bruno Lima", "baixa", 1),
                ],
            )


def get_vectorstore() -> Chroma:
    embeddings = OpenAIEmbeddings(model=OPENAI_EMBEDDING_MODEL)
    return Chroma(
        collection_name="documentos_enviados",
        embedding_function=embeddings,
        persist_directory=str(VECTOR_DIR),
    )


def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def read_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def read_csv(path: Path) -> str:
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as file:
        rows = csv.reader(file)
        return "\n".join(" | ".join(row) for row in rows)


def load_documents_from_inbox() -> list[Document]:
    readers = {
        ".txt": read_text_file,
        ".md": read_text_file,
        ".pdf": read_pdf,
        ".docx": read_docx,
        ".csv": read_csv,
    }
    docs: list[Document] = []
    for path in INBOX_DIR.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in readers:
            continue
        text = readers[path.suffix.lower()](path).strip()
        if text:
            docs.append(Document(page_content=text, metadata={"source": str(path)}))
    return docs


def ingest_uploaded_documents() -> str:
    if DOCUMENT_MEMORY_MODE == "keyword":
        docs = load_documents_from_inbox()
        return (
            f"{len(docs)} documento(s) encontrado(s) em {INBOX_DIR}. "
            "Modo keyword nao precisa criar indice vetorial."
        )

    docs = load_documents_from_inbox()
    if not docs:
        return f"Nenhum documento suportado encontrado em {INBOX_DIR}."

    splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=180)
    chunks = splitter.split_documents(docs)
    vectorstore = get_vectorstore()
    vectorstore.add_documents(chunks)
    return f"{len(docs)} documento(s) indexado(s), gerando {len(chunks)} trecho(s) pesquisaveis."


@tool
def listar_tabelas_do_banco() -> str:
    """Lista as tabelas e colunas disponiveis no banco SQLite local."""
    with sqlite3.connect(DATABASE_PATH) as conn:
        tables = conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table' ORDER BY name"
        ).fetchall()
        lines: list[str] = []
        for (table_name,) in tables:
            if table_name.startswith("sqlite_"):
                continue
            columns = conn.execute(f"PRAGMA table_info({table_name})").fetchall()
            column_text = ", ".join(f"{column[1]} {column[2]}" for column in columns)
            lines.append(f"{table_name}: {column_text}")
        return "\n".join(lines) or "Nenhuma tabela encontrada."


def is_readonly_query(query: str) -> bool:
    normalized = query.strip().lower()
    return normalized.startswith("select ") or normalized.startswith("with ")


@tool
def consultar_banco_sqlite(query: str) -> str:
    """Executa uma consulta SQL somente leitura no banco SQLite local."""
    if not is_readonly_query(query):
        return "Por seguranca, esta ferramenta aceita apenas SELECT ou WITH."

    readonly_uri = f"file:{DATABASE_PATH.as_posix()}?mode=ro"
    try:
        with sqlite3.connect(readonly_uri, uri=True) as conn:
            conn.row_factory = sqlite3.Row
            rows = conn.execute(query).fetchmany(30)
    except sqlite3.Error as exc:
        return f"Erro ao executar SQL: {exc}"

    if not rows:
        return "Consulta executada, sem resultados."

    headers = rows[0].keys()
    lines = [" | ".join(headers)]
    lines.extend(" | ".join(str(row[header]) for header in headers) for row in rows)
    return "\n".join(lines)


@tool
def buscar_documentos_enviados(pergunta: str) -> str:
    """Busca informacoes na memoria criada a partir dos documentos enviados."""
    if DOCUMENT_MEMORY_MODE == "keyword":
        return buscar_documentos_por_palavra_chave(pergunta)

    try:
        vectorstore = get_vectorstore()
        results = vectorstore.similarity_search(pergunta, k=5)
    except Exception as exc:
        return f"Nao consegui buscar nos documentos. Rode /ingest antes. Detalhe: {exc}"

    if not results:
        return "Nenhum trecho relevante encontrado nos documentos enviados."

    return "\n\n".join(format_document_hit(index, doc) for index, doc in enumerate(results, 1))


def format_document_hit(index: int, doc: Document) -> str:
    source = doc.metadata.get("source", "fonte desconhecida")
    snippet = doc.page_content.strip().replace("\n", " ")
    return f"[{index}] Fonte: {source}\n{snippet[:900]}"


def buscar_documentos_por_palavra_chave(pergunta: str) -> str:
    docs = load_documents_from_inbox()
    if not docs:
        return f"Nenhum documento suportado encontrado em {INBOX_DIR}."

    terms = normalized_terms(pergunta)
    scored: list[tuple[int, Document]] = []
    for doc in docs:
        text = doc.page_content.lower()
        score = sum(text.count(term) for term in terms)
        if score > 0:
            scored.append((score, doc))

    if not scored:
        return "Nenhum trecho relevante encontrado nos documentos enviados."

    scored.sort(key=lambda item: item[0], reverse=True)
    return "\n\n".join(
        format_document_hit(index, doc) for index, (_, doc) in enumerate(scored[:5], 1)
    )


def normalized_terms(text: str) -> list[str]:
    terms = re.findall(r"\w+", text.lower(), flags=re.UNICODE)
    ignored = {"a", "o", "e", "de", "da", "do", "das", "dos", "em", "para", "que", "com"}
    return [term for term in terms if len(term) > 2 and term not in ignored]


@tool
def criar_documento_docx(titulo: str, conteudo: str, nome_arquivo: str = "documento.docx") -> str:
    """Cria um documento Word .docx na pasta data/output."""
    safe_name = Path(nome_arquivo).name
    if not safe_name.lower().endswith(".docx"):
        safe_name = f"{safe_name}.docx"

    path = OUTPUT_DIR / safe_name
    doc = DocxDocument()
    doc.add_heading(titulo, level=1)
    for block in split_paragraphs(conteudo):
        doc.add_paragraph(block)
    doc.save(path)
    return f"Documento criado em: {path}"


def split_paragraphs(text: str) -> Iterable[str]:
    paragraphs = [paragraph.strip() for paragraph in text.split("\n\n")]
    return [paragraph for paragraph in paragraphs if paragraph]


def build_agent():
    llm = build_chat_model()
    return create_agent(
        model=llm,
        tools=[
            listar_tabelas_do_banco,
            consultar_banco_sqlite,
            buscar_documentos_enviados,
            criar_documento_docx,
        ],
        system_prompt=(
            "Voce e um agente de automacao em portugues. "
            "Use ferramentas quando precisar consultar banco de dados, pesquisar documentos enviados "
            "ou criar arquivos .docx. Nunca invente dados que deveriam vir do banco ou dos documentos. "
            "Antes de consultar SQL, confira o esquema se necessario. "
            "Mantenha respostas objetivas e explique quais arquivos foram criados."
        ),
    )


def build_chat_model() -> ChatOpenAI:
    if MODEL_PROVIDER == "xai":
        return ChatOpenAI(
            model=XAI_MODEL,
            api_key=os.getenv("XAI_API_KEY"),
            base_url=XAI_BASE_URL,
            temperature=0,
            timeout=3600,
        )
    if MODEL_PROVIDER == "openai":
        return ChatOpenAI(model=OPENAI_MODEL, temperature=0)
    raise ValueError("MODEL_PROVIDER deve ser 'xai' ou 'openai'.")


def missing_api_key_message() -> str | None:
    if MODEL_PROVIDER == "xai" and not os.getenv("XAI_API_KEY"):
        return "Defina XAI_API_KEY no arquivo .env antes de rodar o agente com Grok."
    if MODEL_PROVIDER == "openai" and not os.getenv("OPENAI_API_KEY"):
        return "Defina OPENAI_API_KEY no arquivo .env antes de rodar o agente com OpenAI."
    if DOCUMENT_MEMORY_MODE == "vector" and not os.getenv("OPENAI_API_KEY"):
        return "DOCUMENT_MEMORY_MODE=vector precisa de OPENAI_API_KEY para gerar embeddings."
    return None


def extract_final_text(result: dict) -> str:
    messages = result.get("messages", [])
    for message in reversed(messages):
        content = getattr(message, "content", None)
        if content:
            if isinstance(content, str):
                return content
            return str(content)
    return str(result)


def main() -> None:
    ensure_directories()
    init_sample_database()

    missing_key = missing_api_key_message()
    if missing_key:
        print(missing_key)
        return

    agent = build_agent()
    print("Agente pronto. Use /ingest para indexar data/inbox, ou /sair para encerrar.")

    while True:
        user_input = input("\nvoce> ").strip()
        if not user_input:
            continue
        if user_input.lower() in {"/sair", "sair", "exit", "quit"}:
            print("Encerrando.")
            break
        if user_input.lower() == "/ingest":
            print(ingest_uploaded_documents())
            continue

        result = agent.invoke({"messages": [{"role": "user", "content": user_input}]})
        print(f"\nagente> {extract_final_text(result)}")


if __name__ == "__main__":
    main()
